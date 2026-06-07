#!/usr/bin/env python3
"""Generates genesis-ai-business-plan.docx — a 4-column grid one-pager."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: landscape letter, tight margins ───────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(11)
section.page_height = Inches(8.5)
section.left_margin   = Inches(0.45)
section.right_margin  = Inches(0.45)
section.top_margin    = Inches(0.4)
section.bottom_margin = Inches(0.4)

# ── Colours ───────────────────────────────────────────────────────────────────
BLACK  = RGBColor(0x11, 0x11, 0x11)
YELLOW = RGBColor(0xF5, 0xE6, 0x42)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY  = RGBColor(0x44, 0x44, 0x44)
MGRAY  = RGBColor(0x88, 0x88, 0x88)
LGRAY  = RGBColor(0xF5, 0xF5, 0xF5)
INK    = RGBColor(0x1A, 0x1A, 0x1A)

# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, sides=None):
    """Set cell borders. sides = dict of side: (color, size, val)"""
    if sides is None:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, (color, sz, val) in sides.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def clear_cell(cell):
    """Remove default empty paragraph from cell."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)


def add_heading(cell, text, color=MGRAY, first=True):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    run.font.size  = Pt(6.5)
    run.font.bold  = True
    run.font.color.rgb = color
    run.font.name  = 'Inter'
    return p


def add_para(cell, text, size=8.5, bold=False, color=DGRAY,
             italic=False, space_before=2, space_after=2, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name  = 'Inter'
    return p


def add_bullet(cell, text, size=8.5, color=DGRAY, indent=120):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Pt(8)
    p.paragraph_format.first_line_indent = Pt(-8)
    run = p.add_run(f"•  {text}")
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = 'Inter'
    return p


def add_bold_inline(cell, label, text, size=8.5, label_color=BLACK,
                    text_color=DGRAY, space_before=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(label)
    r1.font.size = Pt(size); r1.font.bold = True
    r1.font.color.rgb = label_color; r1.font.name = 'Inter'
    r2 = p.add_run(text)
    r2.font.size = Pt(size); r2.font.bold = False
    r2.font.color.rgb = text_color; r2.font.name = 'Inter'
    return p


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ════════════════════════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════════════════════════
header_table = doc.add_table(rows=1, cols=2)
header_table.style = 'Table Grid'
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

lc = header_table.cell(0, 0)
rc = header_table.cell(0, 1)
set_cell_border(lc, {
    'top':    ('000000', 0, 'none'),
    'bottom': ('111111', 18, 'single'),
    'left':   ('000000', 0, 'none'),
    'right':  ('000000', 0, 'none'),
})
set_cell_border(rc, {
    'top':    ('000000', 0, 'none'),
    'bottom': ('111111', 18, 'single'),
    'left':   ('000000', 0, 'none'),
    'right':  ('000000', 0, 'none'),
})
set_cell_margins(lc, 0, 60, 0, 0)
set_cell_margins(rc, 0, 60, 0, 0)

# Left: brand name
p = lc.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
r_name = p.add_run("Genesis ⚡")
r_name.font.size = Pt(28); r_name.font.bold = True
r_name.font.color.rgb = BLACK; r_name.font.name = 'Inter'
p2 = lc.add_paragraph()
p2.paragraph_format.space_before = Pt(1)
p2.paragraph_format.space_after  = Pt(0)
r_sub = p2.add_run("INTELLIGENT BRAND DESIGN")
r_sub.font.size = Pt(8); r_sub.font.bold = True
r_sub.font.color.rgb = MGRAY; r_sub.font.name = 'Inter'

# Right: meta
rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = rc.paragraphs[0]
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
r = p.add_run("One-Page Business Plan")
r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLACK; r.font.name = 'Inter'
for line in ["AI-Powered Brand Intelligence Platform  ·  genesisbrands.ai",
             "Darrell Irwin & Saulo Oliveira  ·  June 2026"]:
    p2 = rc.add_paragraph(line)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(8); run.font.color.rgb = MGRAY; run.font.name = 'Inter'

lc.width = Inches(5); rc.width = Inches(5)

doc.add_paragraph()  # small gap
doc.paragraphs[-1].paragraph_format.space_before = Pt(0)
doc.paragraphs[-1].paragraph_format.space_after  = Pt(3)

# ════════════════════════════════════════════════════════════════════════════════
# MAIN GRID TABLE  (4 columns × 3 rows)
# ════════════════════════════════════════════════════════════════════════════════
col_widths = [2.4, 2.35, 2.35, 2.35]

def make_grid(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t

def border_all(cell, color='CCCCCC', sz=4):
    set_cell_border(cell, {s: (color, sz, 'single')
                            for s in ['top','bottom','left','right']})

# ── ROW 1 ─────────────────────────────────────────────────────────────────────
r1 = make_grid(1, 4)
for ci, w in enumerate(col_widths):
    for ri in range(1):
        r1.rows[ri].cells[ci].width = Inches(w)

cells_r1 = [r1.cell(0, i) for i in range(4)]
for c in cells_r1:
    set_cell_margins(c)
    set_cell_bg(c, 'F5F5F5')
    border_all(c)

# Cell 0 — Mission (dark)
c = cells_r1[0]
set_cell_bg(c, '111111')
set_cell_border(c, {s: ('333333', 4, 'single') for s in ['top','bottom','left','right']})
add_heading(c, "Mission", color=YELLOW)
p = c.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(5)
r = p.add_run("The world's first AI creative director for growing businesses.")
r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = 'Inter'
add_para(c, "Genesis AI is an AI-native brand intelligence platform — a trained second brain that conducts discovery, derives strategy, orchestrates specialist agents, and inspects their output through iterative quality loops to deliver complete brand identities in minutes.", size=8, color=RGBColor(0xBB,0xBB,0xBB), space_before=2)
add_para(c, "What makes it unique is not the AI — it is what the AI has been taught. Twenty-five years of brand strategy expertise, encoded into structured reasoning, form the intelligence substrate of the Brain Engine. This is not a general model asked to make logos. It is a purpose-built creative intelligence that took decades to accumulate.", size=8, color=RGBColor(0xBB,0xBB,0xBB), space_before=4)

# Cell 1 — Problem
c = cells_r1[1]
add_heading(c, "Problem")
add_para(c, "Brand is the most critical factor in startup survival — and the least accessible:", size=8.5, color=BLACK, bold=False, space_before=2)
for item in ["5.5M new businesses start in the US every year",
             "Traditional branding costs $15,000–$75,000+",
             "Only ~10% of startups invest in branding in year one",
             "Without a brand, businesses struggle to attract customers, investors, and talent",
             "And even those who buy a logo have no ongoing brand advisor"]:
    add_bullet(c, item, size=8)

# Cell 2 — Differential
c = cells_r1[2]
add_heading(c, "The Differential — The Brain Engine")
add_para(c, "Genesis AI is not a template generator. The core is a trained second brain that:", size=8.5, color=BLACK, space_before=2)
for item in ["Conducts intelligent brand discovery — not a form",
             "Reasons with 25+ years of codified brand strategy",
             "Produces three creative directions simultaneously",
             "Briefs and orchestrates specialist agents",
             "Inspects output in successive rounds until quality threshold is met",
             "Retains brand memory — every session knows your brand",
             "Stays as an ongoing consultant — not a one-time transaction"]:
    add_bullet(c, item, size=8)

# Cell 3 — Customers
c = cells_r1[3]
add_heading(c, "Target Customers")
for item in ["US-based startups and early-stage founders (0–3 years)",
             "Solopreneurs and freelancers building a professional identity",
             "Accelerator and incubator cohorts (B2B channel)",
             "Small businesses needing a rebrand or refresh",
             "University entrepreneurship programs (institutional)"]:
    add_bullet(c, item, size=8)
add_para(c, "Entry through brand creation. Retained through subscription. Expanded through physical products and social media management.", size=7.5, color=MGRAY, italic=True, space_before=5)

# ── ROW 2 ─────────────────────────────────────────────────────────────────────
r2 = make_grid(1, 4)
cells_r2 = [r2.cell(0, i) for i in range(4)]
for ci, (c, w) in enumerate(zip(cells_r2, col_widths)):
    c.width = Inches(w)
    set_cell_margins(c); set_cell_bg(c, 'F5F5F5'); border_all(c)

# Cell 0+1 merged — Revenue Model
cells_r2[0].merge(cells_r2[1])
c = cells_r2[0]
set_cell_bg(c, 'FFFFFF')
add_heading(c, "Revenue Model — Four Layers")

layers = [
    ("Entry — Brand Packages (one-time)  ",  "$229 / $349 / $849",
     "Complete brand identity: logo, copy, colour, typography, brand book, all digital assets. Three creative directions."),
    ("Retention — Subscription  ",            "$39 – $149 / month",
     "Ongoing AI consultant · asset regeneration · social media graphics, carousels & post templates (Growth) · AI social calendar management (Pro)."),
    ("Expansion — Physical Products  ",       "20–40% margin",
     "Branded business cards, banners, t-shirts, mugs, stationery — ordered in-platform, fulfilled by print-on-demand partners. Zero inventory risk."),
    ("Future — AI Social Media Management  ", "Premium tier",
     "The AI manages the brand's social presence: generates, schedules, and publishes content autonomously. Brand voice and identity always consistent."),
]
for i, (title, price, desc) in enumerate(layers):
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3 if i else 2)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Pt(6)
    # Yellow left rule via tab/indent is not easy in docx; use bold label instead
    r1 = p.add_run(title)
    r1.font.size = Pt(8.5); r1.font.bold = True
    r1.font.color.rgb = BLACK if i < 3 else MGRAY; r1.font.name = 'Inter'
    r2 = p.add_run(price)
    r2.font.size = Pt(8.5); r2.font.bold = True
    r2.font.color.rgb = BLACK if i < 3 else MGRAY; r2.font.name = 'Inter'
    p2 = c.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.left_indent = Pt(6)
    run = p2.add_run(desc)
    run.font.size = Pt(7.5)
    run.font.color.rgb = MGRAY if i == 3 else DGRAY
    run.font.italic = i == 3; run.font.name = 'Inter'

# Cell 2 — Competitive Advantage
c = cells_r2[2]
add_heading(c, "Competitive Advantage")
for item in ["Second brain, not templates — reasoning and quality loops",
             "Three directions simultaneously — not one output",
             "25+ years of strategy encoded into AI reasoning",
             "Brand memory compounds with every session",
             "Full loop: creation → subscription → physical → social",
             "No competitor owns this end-to-end journey"]:
    add_bullet(c, item, size=8)
add_para(c, "Tailor Brands and Looka make logos.", size=8, color=MGRAY, italic=True, space_before=5)
add_para(c, "Genesis AI builds brands and stays.", size=8.5, color=BLACK, bold=True, italic=False, space_before=1)

# Cell 3 — Roadmap
c = cells_r2[3]
set_cell_bg(c, 'FFF9CC')
add_heading(c, "Current State & Roadmap")
milestones = [
    ("~5 yrs ago", "Alpha built & validated",  "Decision-tree pipeline, real clients, testimonials on record."),
    ("LIVE NOW",   "Brain Engine operational", "Consultant mode live, delivering brand insights, demo scheduled."),
    ("Next",       "Agent pipeline",            "Connect specialist agents, refine the Brain Engine's prompting & inspection loop."),
    ("Then",       "UI & questionnaire",        "Client-facing discovery experience built and connected."),
    ("EOY 2026",   "Beta launch",               "Full pipeline live, early subscriber access opens."),
]
for when, title, body in milestones:
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{when}  ")
    r1.font.size = Pt(7); r1.font.bold = True
    r1.font.color.rgb = BLACK if when == "LIVE NOW" else MGRAY; r1.font.name = 'Inter'
    r2 = p.add_run(title)
    r2.font.size = Pt(8.5); r2.font.bold = True; r2.font.color.rgb = BLACK; r2.font.name = 'Inter'
    p2 = c.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.left_indent = Pt(8)
    run = p2.add_run(body)
    run.font.size = Pt(7.5); run.font.color.rgb = DGRAY; run.font.name = 'Inter'

# ── ROW 3 ─────────────────────────────────────────────────────────────────────
r3 = make_grid(1, 4)
cells_r3 = [r3.cell(0, i) for i in range(4)]
for ci, (c, w) in enumerate(zip(cells_r3, col_widths)):
    c.width = Inches(w); set_cell_margins(c); set_cell_bg(c, 'F5F5F5'); border_all(c)

# Cell 0+1 merged — Financials
cells_r3[0].merge(cells_r3[1])
c = cells_r3[0]
set_cell_bg(c, 'FFFFFF')
add_heading(c, "3-Year Financial Projections (USD · $349 avg brand package · $59/mo avg subscription)")

# Inner table for projections
proj_table = c._tc  # we'll add a nested table

inner = doc.add_table(rows=1, cols=7)  # temp; we'll move it

# Actually build inside the cell using runs
headers = ["Yr", "Scenario", "Brands/mo", "Packages", "+Subscriptions", "+Physical", "Total"]
data = [
    ("Y1","High","300","$1,258,800","+$382,000","+$72,000","$1,712,800"),
    ("Y1","Mid","150","$629,400","+$191,000","+$36,000","$856,400"),
    ("Y1","Low","75","$314,700","+$95,000","+$18,000","$427,700"),
    ("Y2","High","1,000","$4,188,000","+$2,980,000","+$180,000","$7,348,000"),
    ("Y2","Mid","500","$2,094,000","+$1,490,000","+$90,000","$3,674,000"),
    ("Y2","Low","250","$1,047,000","+$745,000","+$45,000","$1,837,000"),
    ("Y3","High","5,000","$20,940,000","+$8,530,000","+$500,000","$29,970,000"),
    ("Y3","Mid","2,500","$10,470,000","+$4,265,000","+$250,000","$14,985,000"),
    ("Y3","Low","1,250","$5,235,000","+$2,132,000","+$125,000","$7,492,000"),
]
# Header row
p = c.add_paragraph()
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
r = p.add_run("  ".join(f"{h:>12}" if i > 0 else f"{h:<4}" for i,h in enumerate(headers)))
r.font.size = Pt(7); r.font.bold = True; r.font.color.rgb = BLACK; r.font.name = 'Courier New'

yr_colors = {"Y1": False, "Y2": False, "Y3": False}
prev_yr = None
for row in data:
    yr = row[0]
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if yr == prev_yr else 2)
    p.paragraph_format.space_after  = Pt(0)
    line = "  ".join(f"{v:>12}" if i > 0 else f"{v:<4}" for i,v in enumerate(row))
    r = p.add_run(line)
    r.font.size = Pt(7.5)
    r.font.bold = (yr != prev_yr)  # bold for first row of each year
    r.font.color.rgb = BLACK if yr != prev_yr else DGRAY
    r.font.name = 'Courier New'
    prev_yr = yr

add_para(c, "* Seed investment ~$530K · costs grow 20%/yr · subscriptions assume 25% conversion · social management revenue not modelled", size=6.5, color=MGRAY, italic=True, space_before=3)

# Cell 2 — Team
c = cells_r3[2]
add_heading(c, "Team")
add_para(c, "Darrell Irwin — Founder & CEO", size=9, bold=True, color=BLACK, space_before=2)
add_para(c, "25+ years brand strategy. Founded Cre8ion agency. The methodology and intelligence powering the Brain Engine is his life's work, codified into structured AI reasoning.", size=8, color=DGRAY, space_before=1)
add_para(c, "Saulo Oliveira — Co-Founder & CTO", size=9, bold=True, color=BLACK, space_before=6)
add_para(c, "Enterprise software architect & AI systems engineer. Designed the multi-agent orchestration engine, the Brain Engine, and the two-layer knowledge infrastructure.", size=8, color=DGRAY, space_before=1)
add_para(c, "$300,000+ in time invested to date.", size=9, bold=True, color=BLACK, space_before=6)

# Cell 3 — Go-to-Market
c = cells_r3[3]
add_heading(c, "Go-to-Market (US)")
add_bold_inline(c, "Partnerships:  ", "SBA programs, YCombinator, Techstars, 500 Startups, university incubators", size=8, space_before=2)
add_bold_inline(c, "Content:  ", "Founder personal brand + YouTube brand-building series", size=8, space_before=2)
add_bold_inline(c, "Performance:  ", "Google & Meta targeting founders & entrepreneurs", size=8, space_before=2)
add_bold_inline(c, "B2B:  ", "White-label licensing to accelerators, coworking hubs, and incubators", size=8, space_before=2)
add_para(c, "Market", size=8, bold=True, color=BLACK, space_before=6)
add_para(c, "TAM  5.5M new US businesses/yr\nSAM  1.5M serviceable\nSOM  150K obtainable (10%)", size=8, color=DGRAY, space_before=1)
add_para(c, "Seed investment sought: $530,000", size=8.5, bold=True, color=BLACK, space_before=5)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.paragraphs[-1]
fp.paragraph_format.space_before = Pt(3)
fp.paragraph_format.space_after  = Pt(0)
fp.add_run("Genesis AI  ·  Intelligent Brand Design  ·  genesisbrands.ai     |     ")
fp.add_run("Darrell Irwin, Founder & CEO  ·  Saulo Oliveira, Co-Founder & CTO")
for run in fp.runs:
    run.font.size = Pt(7.5); run.font.color.rgb = MGRAY; run.font.name = 'Inter'

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/saulo/genesis-ai/docs/eb2-niw/genesis-ai-business-plan.docx"
doc.save(out)
print(f"Saved: {out}")
